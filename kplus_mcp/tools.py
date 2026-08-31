"""Единый реестр инструментов К+.

Это источник правды. MCP-сервер, CLI-агент и HTTP-фасад ничего своего
не реализуют — все три просто оборачивают этот список. Добавили инструмент
здесь — он сразу доступен во всех трёх режимах.
"""
from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Awaitable, Callable

from . import browser, config, pace
from .errors import log_event, log_exception
from .pace import PaceLimit


def _j(obj: Any) -> str:
    return json.dumps(obj, ensure_ascii=False, indent=2)


async def _looks_logged_in() -> bool:
    p = await browser.page()
    html = (await p.content()).lower()
    has_password = await p.locator("input[type=password]").count() > 0
    return not has_password and "войти" not in html[:4000]


# --- реализации ----------------------------------------------------------------

async def kplus_status() -> str:
    p = await browser.page()
    if not p.url.startswith("http"):
        await browser.goto(config.BASE_URL)
        p = await browser.page()
    return _j({
        "url": p.url,
        "title": await p.title(),
        "logged_in": await _looks_logged_in(),
        "base_url": config.BASE_URL,
        "profile": str(config.PROFILE_DIR),
    })


async def kplus_login() -> str:
    p = await browser.goto(config.BASE_URL)
    if await _looks_logged_in():
        return "Сессия уже активна: " + p.url

    login = os.environ.get("KPLUS_LOGIN")
    password = os.environ.get("KPLUS_PASSWORD")
    if login and password:
        try:
            await p.locator("input[type=text], input[type=email], input[name*=login]").first.fill(login)
            await p.locator("input[type=password]").first.fill(password)
        except Exception:
            pass

    waited, step = 0, 3
    while waited < config.LOGIN_WAIT_S:
        await p.wait_for_timeout(step * 1000)
        waited += step
        if await _looks_logged_in():
            return f"Вход выполнен за ~{waited} с. Текущий адрес: {p.url}"
    return ("Вход не подтверждён за отведённое время. Окно браузера осталось открытым — "
            "завершите вход вручную и вызовите kplus_status.")


async def kplus_goto(url: str) -> str:
    await browser.goto(url)
    return _j(await browser.snapshot())


async def kplus_snapshot(text_limit: int = 0) -> str:
    return _j(await browser.snapshot(text_limit or None))


async def kplus_click(ref: str) -> str:
    await browser.click(ref)
    return _j(await browser.snapshot())


async def kplus_fill(ref: str, value: str, submit: bool = True) -> str:
    await browser.fill(ref, value, submit=submit)
    return _j(await browser.snapshot())


async def kplus_search(query: str) -> str:
    p = await browser.page()
    selectors = [
        "input[type=search]",
        'input[placeholder*="оиск"]',
        'input[name*="search"]',
        'textarea[placeholder*="оиск"]',
        "#searchInput",
        "input[type=text]",
    ]
    for attempt in range(2):
        for sel in selectors:
            loc = p.locator(sel).first
            try:
                if await loc.count() and await loc.is_visible():
                    await loc.click()
                    await loc.fill(query)
                    await pace.before_request()
                    await loc.press("Enter")
                    await browser.settle(p, 2000)
                    return _j(await browser.snapshot())
            except Exception:
                continue
        if attempt == 0:
            p = await browser.goto(config.BASE_URL)
    return ("Поле поиска не найдено. Сделайте kplus_snapshot и кликните нужный элемент "
            "через kplus_click вручную.")


async def kplus_read_document(offset: int = 0) -> str:
    text = await browser.document_text()
    chunk = text[offset: offset + config.DOC_CHUNK]
    return _j({
        "total_chars": len(text),
        "offset": offset,
        "next_offset": offset + len(chunk) if offset + len(chunk) < len(text) else None,
        "text": chunk,
    })


async def kplus_save(filename: str, content: str, docx: bool = True) -> str:
    config.ensure_dirs()
    stem = re.sub(r"[^\w\-. ]+", "_", filename).strip() or "document"
    stem = re.sub(r"\.(md|docx)$", "", stem, flags=re.I)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M")
    md_path = config.OUT_DIR / f"{stamp}_{stem}.md"
    md_path.write_text(content, encoding="utf-8")
    written = [str(md_path)]

    if docx:
        try:
            from docx import Document

            doc = Document()
            for line in content.splitlines():
                s = line.rstrip()
                if s.startswith("### "):
                    doc.add_heading(s[4:], level=3)
                elif s.startswith("## "):
                    doc.add_heading(s[3:], level=2)
                elif s.startswith("# "):
                    doc.add_heading(s[2:], level=1)
                elif s.startswith(("- ", "* ")):
                    doc.add_paragraph(s[2:], style="List Bullet")
                else:
                    doc.add_paragraph(s)
            docx_path = config.OUT_DIR / f"{stamp}_{stem}.docx"
            doc.save(str(docx_path))
            written.append(str(docx_path))
        except ImportError:
            written.append("(python-docx не установлен — .docx пропущен)")

    return "Сохранено:\n" + "\n".join(written)


async def kplus_close() -> str:
    """Закрыть окно и тем самым освободить сеанс К+.

    Именно закрыть, а не отключиться: пока окно живо, К+ считает сеанс
    занятым, и вход с другого устройства будет выглядеть как второй
    пользователь.
    """
    await browser.shutdown()
    browser.quit_browser()
    return "Окно К+ закрыто, сеанс освобождён."


# --- реестр ---------------------------------------------------------------------

@dataclass(frozen=True)
class Tool:
    name: str
    description: str
    parameters: dict
    fn: Callable[..., Awaitable[str]]


def _schema(props: dict, required: list[str] | None = None) -> dict:
    return {"type": "object", "properties": props, "required": required or []}


TOOLS: list[Tool] = [
    Tool("kplus_status",
         "Состояние: открыт ли браузер, авторизована ли сессия К+, текущий адрес. "
         "Всегда вызывайте первым.",
         _schema({}), kplus_status),
    Tool("kplus_login",
         "Открыть окно К+ и дождаться входа пользователя (до 5 минут). "
         "Логин и пароль вводит человек в окне браузера. Сессия сохраняется "
         "в профиле. Работать строго в этом одном окне: КонсультантПлюс "
         "разрешает один сеанс и считает второе окно вторым пользователем.",
         _schema({}), kplus_login),
    Tool("kplus_goto",
         "Перейти по адресу внутри разрешённого домена К+. Возвращает снимок страницы.",
         _schema({"url": {"type": "string", "description": "Полный URL на сервере К+"}},
                 ["url"]), kplus_goto),
    Tool("kplus_search",
         "Быстрый поиск по К+: найти строку поиска, ввести запрос, отправить. "
         "Возвращает снимок страницы с результатами.",
         _schema({"query": {"type": "string", "description": "Поисковый запрос на русском"}},
                 ["query"]), kplus_search),
    Tool("kplus_snapshot",
         "Снимок текущей страницы: адрес, заголовок, видимый текст и список "
         "кликабельных элементов с ref-идентификаторами.",
         _schema({"text_limit": {"type": "integer",
                                 "description": "Лимит символов текста, 0 — по умолчанию"}}),
         kplus_snapshot),
    Tool("kplus_click",
         "Кликнуть элемент по ref из последнего снимка. Возвращает новый снимок.",
         _schema({"ref": {"type": "string", "description": "ref вида e12"}}, ["ref"]), kplus_click),
    Tool("kplus_fill",
         "Ввести текст в поле по ref. По умолчанию сразу отправляет Enter.",
         _schema({"ref": {"type": "string"}, "value": {"type": "string"},
                  "submit": {"type": "boolean"}}, ["ref", "value"]), kplus_fill),
    Tool("kplus_read_document",
         "Вычитать текст открытого документа порциями. offset — смещение в символах. "
         "В ответе поле next_offset, если есть продолжение.",
         _schema({"offset": {"type": "integer"}}), kplus_read_document),
    Tool("kplus_save",
         "Сохранить готовый документ в папку вывода. content — Markdown. "
         "Пишется .md и .docx.",
         _schema({"filename": {"type": "string"}, "content": {"type": "string"},
                  "docx": {"type": "boolean"}}, ["filename", "content"]), kplus_save),
    Tool("kplus_close",
         "Закрыть окно К+ и освободить сеанс. Вызывать, когда работа закончена: "
         "пока окно открыто, К+ считает сеанс занятым.",
         _schema({}), kplus_close),
]

BY_NAME: dict[str, Tool] = {t.name: t for t in TOOLS}


async def call(name: str, args: dict | None = None) -> str:
    tool = BY_NAME.get(name)
    if tool is None:
        return f"Неизвестный инструмент: {name}. Доступны: {', '.join(BY_NAME)}"
    log_event(f"Инструмент {name} {json.dumps(args or {}, ensure_ascii=False)[:200]}")
    try:
        return await tool.fn(**(args or {}))
    except PaceLimit as e:
        # Не поломка, а сработавший ограничитель нагрузки. Модель должна
        # увидеть причину и закончить работу тем, что уже собрано.
        log_event(f"Ограничитель темпа остановил {name}")
        return (f"ОСТАНОВЛЕНО ОГРАНИЧИТЕЛЕМ: {e}\n"
                "Больше страниц не открывайте. Подведите итог по тому, "
                "что уже прочитано, и честно скажите, чего не хватило.")
    except TypeError as e:
        return f"Неверные аргументы для {name}: {e}"
    except Exception as e:
        # Модель должна увидеть ошибку и попробовать иначе, а не упасть.
        # Полная трассировка при этом уходит в журнал для разбора.
        log_exception(f"инструмент {name}", e)
        return f"Ошибка {name}: {type(e).__name__}: {e}"


def openai_schema() -> list[dict]:
    """Описание инструментов в формате OpenAI-совместимого function calling."""
    return [{"type": "function",
             "function": {"name": t.name, "description": t.description,
                          "parameters": t.parameters}} for t in TOOLS]
